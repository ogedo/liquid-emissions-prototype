"""
Generate: Liquid Emissions Revenue Collection — Software Requirements Specification (SRS)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x3E)
TEAL   = RGBColor(0x00, 0x8B, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xF8)
DARK_T = RGBColor(0x1A, 0x2E, 0x4A)
AMBER  = RGBColor(0xD4, 0x7E, 0x00)
GREY   = RGBColor(0x80, 0x80, 0x80)

# ── Helpers ───────────────────────────────────────────────────────────────────

def rgb_hex(rgb):
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        run = p.add_run(
            'Liquid Emissions Revenue Collection  |  Software Requirements Specification v1.0  |  '
            'Confidential — Internal Use Only'
        )
        run.font.size = Pt(8)
        run.font.color.rgb = GREY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = NAVY
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), rgb_hex(TEAL))
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = TEAL
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = DARK_T
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    elif level == 4:
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)

def add_body(doc, text, italic=False, bullet=False, size=10):
    style = 'List Bullet' if bullet else 'Normal'
    p = doc.add_paragraph(style=style)
    p.clear()
    if bullet:
        # re-add bullet marker manually for plain style
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f'\u2022  {text}')
    else:
        run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.clear()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'NOTE: {text}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = AMBER

def make_table(doc, headers, rows, col_widths=None, compact=False):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hrow = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, DARK_T)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(hdr)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(8 if compact else 9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
    # data rows
    for idx, row in enumerate(rows):
        drow = tbl.add_row()
        for i, val in enumerate(row):
            cell = drow.cells[i]
            if idx % 2 == 1:
                set_cell_bg(cell, LIGHT)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.size = Pt(8 if compact else 9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()
    return tbl

def req_table(doc, rows):
    """Standard requirement table: ID | Description | Priority | Source"""
    return make_table(
        doc,
        ['Req ID', 'Requirement Description', 'Priority', 'Notes / Source'],
        rows,
        col_widths=[0.85, 4.0, 0.7, 1.45],
        compact=True,
    )

def spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

# ════════════════════════════════════════════════════════════════════════════
# BUILD
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
add_footer(doc)

# ── COVER ────────────────────────────────────────────────────────────────────
spacer(doc, 4)
for txt, sz, bold, color in [
    ('Liquid Emissions Revenue Collection',     28, True,  NAVY),
    ('Software Requirements Specification',     20, True,  TEAL),
    ('SRS — Version 1.0',                       12, False, RGBColor(0x44,0x44,0x44)),
    ('',                                         8, False, WHITE),
    ('Prepared by: MetronBase Labs',            10, False, DARK_T),
    ('Classification: Confidential — Internal', 9,  True,  AMBER),
    ('Date: March 2026',                         9, False, GREY),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ── 1. INTRODUCTION ──────────────────────────────────────────────────────────
add_heading(doc, '1. Introduction')

add_heading(doc, '1.1  Purpose', 2)
add_body(doc,
    'This Software Requirements Specification (SRS) defines the functional and non-functional '
    'requirements for the Liquid Emissions Revenue Collection platform — a multi-tenant, '
    'multi-currency digital system that automates the collection, settlement, and reporting of '
    'ship emission fees at Nigerian ports.')
add_body(doc,
    'The platform is operated by Blue Wave on behalf of NIMASA, and developed by MetronBase Labs / '
    'Liquid Payments. This document is the authoritative technical reference for all development '
    'phases and must be read alongside the Domain Model and Prisma Schema.')

add_heading(doc, '1.2  Scope', 2)
add_body(doc, 'The system covers:')
for item in [
    'Vessel operator account creation and identity management',
    'Retrieval of emission assessment data from an external maritime authority system via Payment Reference ID',
    'Collection of emission fees plus platform service charges in USD and NGN (crypto/stablecoin in a future phase)',
    'Multi-party settlement disbursement: NIMASA, Port Authority, Blue Wave, Liquid Payments',
    'Payouts to Nigerian bank accounts (NGN) and international accounts via Wise Business API (USD)',
    'Audit trail, revenue reporting, and compliance exports',
    'Notification delivery (email, in-app; SMS optional)',
]:
    add_body(doc, item, bullet=True)
add_body(doc, 'Out of scope for v1.0: vessel AIS/SEMSW sensor integration, own payment gateway, crypto collection.')

add_heading(doc, '1.3  Definitions & Acronyms', 2)
defs = [
    ('AIS',             'Automatic Identification System — vessel tracking'),
    ('BullMQ',          'Redis-backed job queue library'),
    ('CO2e',            'Carbon dioxide equivalent — standardised emission unit'),
    ('DPIA',            'Data Privacy Impact Assessment'),
    ('FX',              'Foreign Exchange rate'),
    ('GRT',             'Gross Registered Tonnage'),
    ('IMO',             'International Maritime Organization vessel number'),
    ('MARPOL',          'International Convention for the Prevention of Pollution from Ships'),
    ('NDPR',            'Nigeria Data Protection Regulation'),
    ('NIMASA',          'Nigerian Maritime Administration and Safety Agency'),
    ('Payment Ref ID',  'Unique reference issued by the external emission authority to vessel operators'),
    ('PCI-DSS',         'Payment Card Industry Data Security Standard'),
    ('PRD',             'Product Requirements Document'),
    ('RBAC',            'Role-Based Access Control'),
    ('SEMSW',           'Ship Emission Monitoring System with Wireless — physical sensor system'),
    ('SRS',             'This document — Software Requirements Specification'),
    ('TOTP',            'Time-based One-Time Password (2FA)'),
    ('USDT / USDC',     'USD-pegged stablecoins (Tether / Circle)'),
]
make_table(doc, ['Term / Acronym', 'Definition'], defs, col_widths=[1.6, 5.4])

add_heading(doc, '1.4  References', 2)
for ref in [
    'Liquid Emissions Revenue Collection.docx — Product Requirements Document (PRD)',
    'Liquid_Emissions_Domain_Model (1).pdf — DDD Strategic & Tactical Domain Model',
    'Liquid_Emissions_Prisma_Schema.docx — Full Prisma Database Schema',
    'Emmission Fees (1).pdf — Emission Fee Comparison: EU ETS / US / Nigeria',
    'emissions implementation (1).pdf — SEMSW Physical Monitoring Technology',
    'MetronBase_Labs_Playbook.docx — Agency Process & Document Framework',
    'MARPOL Annex VI — IMO Emission Standards Reference',
    'NDPR 2019 / GDPR 2018 — Data Protection Compliance Frameworks',
]:
    add_body(doc, ref, bullet=True)

# ── 2. OVERALL DESCRIPTION ───────────────────────────────────────────────────
add_heading(doc, '2. Overall Description')

add_heading(doc, '2.1  Product Perspective', 2)
add_body(doc,
    'The Liquid Emissions platform is a standalone web application that sits at the intersection '
    'of maritime compliance, fintech, and government revenue collection. It does not generate '
    'emission assessments — it consumes them from an authoritative external system and provides '
    'the payment, settlement, and reporting infrastructure on top.')
add_body(doc,
    'The system integrates with: (a) an external maritime emission authority API, (b) payment '
    'gateways (Paystack, Flutterwave), (c) a payout provider (Wise Business API for USD; '
    'Paystack/Flutterwave Transfer API for NGN), and (d) an email/SMS notification provider.')

add_heading(doc, '2.2  Core User Flow (Happy Path)', 2)
steps = [
    ('1', 'External emission authority assesses a vessel voyage and generates an emission record with a unique Payment Reference ID'),
    ('2', 'Port authority / NIMASA notifies the vessel operator (by email or physical document) of their Payment Reference ID and amount due'),
    ('3', 'Vessel operator logs in to the platform (account required — no guest access)'),
    ('4', 'Operator enters their Payment Reference ID'),
    ('5', 'Platform calls the external system API, retrieves assessment data (fee amount, vessel info, voyage details)'),
    ('6', 'Platform creates a PaymentOrder: base emission fee + platform service charge, denominated in USD'),
    ('7', 'Operator selects payment currency (USD or NGN); platform locks FX rate for NGN at that moment'),
    ('8', 'Payment processed via Paystack or Flutterwave; webhook confirms completion'),
    ('9', 'Platform triggers settlement: wallets credited per SplitRule (NIMASA, Port Authority, Blue Wave, Liquid Payments)'),
    ('10', 'Receipt emailed to operator; domain events recorded to outbox; notification dispatched to external system'),
]
make_table(doc, ['Step', 'Action'], steps, col_widths=[0.4, 6.6])

add_heading(doc, '2.3  User Classes & Characteristics', 2)
users = [
    ('Vessel Operator',  'OPERATOR',  'Pays emission fees. Manages their organisation\'s vessels. Views payment history and receipts.'),
    ('Finance Officer',  'FINANCE',   'Views settlement statements, wallet balances, payout history for their organisation.'),
    ('Platform Admin',   'ADMIN',     'Manages organisations, users, fee schedules, split rules, gateways. Requires 2FA.'),
    ('Auditor',          'AUDITOR',   'Read-only access across all data. Cannot modify any record.'),
    ('System / Webhook', 'SYSTEM',    'Internal service account for gateway webhooks and external API callbacks.'),
]
make_table(doc, ['User Class', 'Role', 'Characteristics'], users, col_widths=[1.4, 0.9, 4.7])

add_heading(doc, '2.4  Operating Environment', 2)
env = [
    ('Cloud Provider',      'Railway (PostgreSQL managed) + any VPS/cloud for API containers'),
    ('Frontend',            'Next.js 14 — deployed to Vercel or equivalent CDN-backed host'),
    ('Backend',             'Node.js + Express — containerised, horizontally scalable'),
    ('Database',            'PostgreSQL 15 (Railway managed)'),
    ('Cache / Queue',       'Redis (Upstash or self-hosted) + BullMQ'),
    ('Payment Gateways',    'Paystack (NGN primary), Flutterwave (USD + NGN)'),
    ('USD Payout',          'Wise Business API (primary); Flutterwave wire (fallback)'),
    ('NGN Payout',          'Paystack Transfer API + Flutterwave Transfer API'),
    ('Notification',        'SendGrid / Mailgun (email); Termii / Africa\'s Talking (SMS)'),
    ('Browser Support',     'Chrome, Firefox, Edge, Safari — latest two major versions'),
    ('Mobile',              'Responsive web; no native mobile app in v1.0'),
]
make_table(doc, ['Environment Component', 'Detail'], env, col_widths=[2.0, 5.0])

add_heading(doc, '2.5  Assumptions & Dependencies', 2)
assumptions = [
    'The external maritime emission authority exposes a REST or SOAP API that accepts a Payment Reference ID and returns structured assessment data.',
    'The external API returns either a pre-calculated fee amount in USD, or sufficient raw inputs (fuel consumed, emission factor) for the platform to calculate fee independently.',
    'All vessels subject to fees are ≥ 5,000 GRT, per MARPOL and NIMASA regulatory scope.',
    'The fixed base emission fee is USD 100 per ton CO2e; the platform service charge is configured separately in the FeeSchedule.',
    'Paystack supports USD card collection for international vessel operators.',
    'Flutterwave provides USD collection for international payments.',
    'Wise Business API is available for the platform\'s operational entity and supports NGN → USD and USD payout flows.',
    'NIMASA and Port Authority payout bank accounts are pre-registered in the system before settlement disbursements begin.',
    'The platform will not store raw card data; PCI-DSS scope is SAQ-A (fully hosted gateway pages).',
    'FX rates for NGN/USD are fetched from a live rate provider at payment initiation and locked for the life of the PaymentOrder.',
]
for a in assumptions:
    add_body(doc, a, bullet=True)

# ── 3. EXTERNAL SYSTEM INTEGRATION ───────────────────────────────────────────
add_heading(doc, '3. External System Integration — Emission Authority API')
add_body(doc,
    'The emission data pipeline is the most critical external dependency. The platform does not '
    'generate emission assessments — it acts as the collection layer on top of the maritime '
    'authority\'s assessment engine. This section defines the integration contract.')

add_heading(doc, '3.1  Payment Reference ID Flow', 2)
add_body(doc,
    'The Payment Reference ID (hereafter "Payment Ref ID") is a unique, opaque token or '
    'structured reference number generated by the external maritime authority system when a '
    'vessel\'s emission assessment is finalised. It is the primary key linking the platform\'s '
    'PaymentOrder to the authority\'s assessment record.')

add_heading(doc, '3.2  External API Contract (Required Capabilities)', 2)
api_caps = [
    ('GET /assessment/{paymentRefId}', 'Fetch assessment data by Payment Ref ID', 'REQUIRED'),
    ('POST /payment-confirmation',     'Notify external system of successful payment; send receipt reference', 'REQUIRED'),
    ('GET /assessment/status/{id}',    'Check whether an assessment has already been paid (idempotency guard)', 'REQUIRED'),
    ('GET /vessel/{imoNumber}',        'Fetch vessel profile by IMO number for display/validation', 'OPTIONAL'),
    ('Webhook (inbound)',              'External system may push updated assessment data to platform', 'FUTURE'),
]
make_table(doc, ['Endpoint / Capability', 'Purpose', 'v1.0 Status'],
           api_caps, col_widths=[2.2, 3.4, 1.4])

add_heading(doc, '3.3  Assessment Response Payload (Expected Fields)', 2)
payload = [
    ('paymentRefId',        'string',  'REQUIRED', 'Unique reference — matches the input'),
    ('imoNumber',           'string',  'REQUIRED', 'IMO vessel identifier'),
    ('vesselName',          'string',  'REQUIRED', 'Vessel display name'),
    ('vesselGRT',           'integer', 'REQUIRED', 'Gross Registered Tonnage'),
    ('portOfCall',          'string',  'REQUIRED', 'UN/LOCODE of Nigerian port'),
    ('voyageArrival',       'date',    'REQUIRED', 'ISO 8601 date'),
    ('voyageDeparture',     'date',    'REQUIRED', 'ISO 8601 date'),
    ('feeAmountUSD',        'integer', 'PREFERRED','Pre-calculated fee in USD minor units (cents). If absent, platform calculates.'),
    ('co2eTons',            'decimal', 'PREFERRED','CO2e quantity. Used if feeAmountUSD absent.'),
    ('fuelConsumedMT',      'decimal', 'OPTIONAL', 'Metric tons of fuel consumed. Used for independent calculation.'),
    ('emissionFactorKgPerMT','decimal','OPTIONAL', 'Emission factor — used with fuelConsumedMT if co2eTons absent.'),
    ('feeScheduleRef',      'string',  'OPTIONAL', 'Authority\'s fee schedule version reference'),
    ('assessmentStatus',    'enum',    'REQUIRED', 'PENDING_PAYMENT | PAID | CANCELLED | DISPUTED'),
    ('currency',            'string',  'REQUIRED', 'Always "USD" — authority bills in USD'),
]
make_table(doc, ['Field', 'Type', 'Requirement', 'Description'],
           payload, col_widths=[1.6, 0.7, 0.9, 3.8], compact=True)

add_note(doc,
    'If feeAmountUSD is provided by the external system, the platform uses it directly. '
    'If absent, the platform computes: CO2e_tons × USD 100. If neither is provided, '
    'the PaymentOrder is held in ASSESSMENT_INCOMPLETE status and an alert is raised.')

add_heading(doc, '3.4  Platform Charges', 2)
add_body(doc,
    'In addition to the base emission fee received from the external system, the platform '
    'appends a service charge. This charge is configured in the FeeSchedule aggregate and '
    'may be expressed as:')
for item in [
    'A flat USD amount per payment (e.g., USD 50.00)',
    'A percentage of the base fee (e.g., 1.5%)',
    'A combination: flat + percentage',
]:
    add_body(doc, item, bullet=True)
add_body(doc,
    'The total amount presented to the vessel operator is: '
    'Base Fee (from external system) + Platform Service Charge. Both line items are shown '
    'separately on the payment screen and on the receipt.')

# ── 4. FUNCTIONAL REQUIREMENTS ───────────────────────────────────────────────
add_heading(doc, '4. Functional Requirements')
add_body(doc,
    'Requirements are grouped by bounded context. Priority: '
    'MUST = MVP required | SHOULD = launch required | COULD = post-launch | WON\'T = explicitly out of scope.')

# 4.1 Identity
add_heading(doc, '4.1  Identity & Access Management (IDN)', 2)
req_table(doc, [
    ('IDN-01', 'Account creation is mandatory. No guest or anonymous payment access.', 'MUST', 'Business decision'),
    ('IDN-02', 'An Organisation record must be created before any user can make a payment. One user may belong to exactly one Organisation.', 'MUST', 'Domain invariant'),
    ('IDN-03', 'Supported roles: ADMIN, OPERATOR, FINANCE, AUDITOR. Roles are assigned per user within an Organisation.', 'MUST', 'RBAC'),
    ('IDN-04', 'AUDITOR role grants read-only access at the domain level. No create, update, or delete operations permitted.', 'MUST', 'Compliance'),
    ('IDN-05', 'ADMIN role requires TOTP-based 2FA (e.g., Google Authenticator). Login without 2FA is rejected.', 'MUST', 'Security'),
    ('IDN-06', 'Authentication uses JWT (RS256 asymmetric signing). Access tokens expire in 15 minutes; refresh tokens in 7 days.', 'MUST', 'Security'),
    ('IDN-07', 'An Organisation may register one or more vessel IMO numbers. A vessel IMO may only be active under one Organisation at a time.', 'MUST', 'Domain integrity'),
    ('IDN-08', 'Password reset via time-limited email token (expires 1 hour). Brute-force protection: 5 failed attempts triggers 15-minute lockout.', 'MUST', 'Security'),
    ('IDN-09', 'All user actions are logged with user ID, timestamp, IP address, and action type.', 'MUST', 'Audit / NDPR'),
    ('IDN-10', 'An Organisation may be suspended by ADMIN, preventing all further payments from that org.', 'SHOULD', 'Operations'),
    ('IDN-11', 'SSO (SAML 2.0 / OAuth2) for enterprise organisations.', 'COULD', 'Future phase'),
])

# 4.2 Emission Assessment Ingestion
add_heading(doc, '4.2  Emission Assessment Ingestion (EMS)', 2)
req_table(doc, [
    ('EMS-01', 'A vessel operator enters a Payment Reference ID on the platform to initiate a payment.', 'MUST', 'Core flow'),
    ('EMS-02', 'On receipt of a Payment Ref ID, the system calls the external emission authority API to retrieve assessment data.', 'MUST', 'External API'),
    ('EMS-03', 'The system validates that assessmentStatus == PENDING_PAYMENT before creating a PaymentOrder. Any other status (PAID, CANCELLED, DISPUTED) is rejected with a clear error message.', 'MUST', 'Idempotency guard'),
    ('EMS-04', 'Assessment data retrieved from the external system is stored as an EmissionAssessment record, linked to the PaymentOrder by paymentRefId.', 'MUST', 'Data persistence'),
    ('EMS-05', 'If the external system provides feeAmountUSD, the platform uses it as the base fee without modification.', 'MUST', 'Fee fidelity'),
    ('EMS-06', 'If feeAmountUSD is absent but co2eTons is provided, the platform calculates: baseFeeCents = co2eTons × 100 × 100 (USD minor units).', 'MUST', 'Fallback calculation'),
    ('EMS-07', 'If neither feeAmountUSD nor co2eTons is provided, the PaymentOrder is created in ASSESSMENT_INCOMPLETE status and an admin alert is triggered. Payment cannot proceed until resolved.', 'MUST', 'Error handling'),
    ('EMS-08', 'The platform service charge is appended to the base fee. Both amounts are line-itemised on the PaymentOrder and visible to the operator before payment.', 'MUST', 'Transparency'),
    ('EMS-09', 'One Payment Ref ID may only be paid once. A second attempt to use a paid reference is rejected (check: assessmentStatus == PAID on external system + local DB check).', 'MUST', 'Idempotency'),
    ('EMS-10', 'On successful payment, the platform notifies the external system via POST /payment-confirmation with: paymentRefId, platformPaymentRef, timestamp, amountPaidUSD.', 'MUST', 'External notification'),
    ('EMS-11', 'If the external API is unavailable at lookup time, the system returns an error to the user and does not create a PaymentOrder. Retry is manual.', 'MUST', 'Resilience'),
    ('EMS-12', 'The system caches assessment data per Payment Ref ID for 15 minutes (Redis) to avoid redundant external API calls during the checkout session.', 'SHOULD', 'Performance'),
    ('EMS-13', 'Admin can manually override an EmissionAssessment (e.g., dispute resolution) with a mandatory audit note.', 'SHOULD', 'Operations'),
])

# 4.3 Payment Processing
add_heading(doc, '4.3  Payment Processing (PAY)', 2)
req_table(doc, [
    ('PAY-01', 'Billing currency is always USD. All PaymentOrder amounts are stored and reported in USD minor units (cents).', 'MUST', 'Currency standard'),
    ('PAY-02', 'The operator may choose to pay in USD or NGN. Currency selection is presented after assessment data is loaded.', 'MUST', 'Multi-currency'),
    ('PAY-03', 'For NGN payments, the platform fetches the live USD/NGN exchange rate from a configured FX provider at payment initiation. The rate is locked on the PaymentOrder for its lifetime.', 'MUST', 'FX handling'),
    ('PAY-04', 'The NGN equivalent is calculated as: ngnAmount = (usdAmountCents / 100) × fxRate × 100 (in kobo minor units). Both USD and NGN amounts are stored on the PaymentOrder.', 'MUST', 'Data integrity'),
    ('PAY-05', 'USD payments are processed via Flutterwave (primary) or Paystack (secondary) USD collection.', 'MUST', 'USD gateway'),
    ('PAY-06', 'NGN payments are processed via Paystack (primary) or Flutterwave (secondary).', 'MUST', 'NGN gateway'),
    ('PAY-07', 'The gateway routing is configurable by ADMIN without code deployment (gateway priority config in DB).', 'SHOULD', 'Ops flexibility'),
    ('PAY-08', 'Crypto (USDT, USDC) payment collection is out of scope for v1.0 but the gateway adapter interface MUST be designed to support plugging in a crypto gateway in a future phase.', 'MUST', 'Architecture constraint'),
    ('PAY-09', 'PaymentOrder state machine: PENDING → PROCESSING → COMPLETED | FAILED | REFUNDED. No other transitions are valid.', 'MUST', 'Domain invariant'),
    ('PAY-10', 'A completed PaymentOrder is immutable. Refunds are processed as a new reversal PaymentOrder referencing the original.', 'MUST', 'Domain invariant'),
    ('PAY-11', 'The TransactionReference field is globally unique across all PaymentOrders. Generated as a platform-scoped UUID + timestamp hash.', 'MUST', 'Uniqueness'),
    ('PAY-12', 'Idempotency: a Redis key (TTL 24 hours) keyed on {paymentRefId}:{currency} prevents duplicate PaymentOrder creation from concurrent requests.', 'MUST', 'Idempotency'),
    ('PAY-13', 'Gateway webhook payloads are verified via HMAC signature before processing. Unverified webhooks are rejected and logged.', 'MUST', 'Security'),
    ('PAY-14', 'Webhook events are processed via BullMQ jobs (not synchronously in the webhook handler) to ensure durability and retry on failure.', 'MUST', 'Reliability'),
    ('PAY-15', 'Own gateway: the gateway integration MUST use an adapter/port pattern (interface + implementation). A future own-gateway implementation requires only a new adapter — no core payment logic changes.', 'MUST', 'Architecture'),
    ('PAY-16', 'Payment link generation: Admin may generate a one-time payment link for a specific Payment Ref ID to share with operators who cannot self-serve.', 'SHOULD', 'Ops'),
    ('PAY-17', 'Partial payments are not supported. A PaymentOrder must be fully settled in a single transaction.', 'MUST', 'Business rule'),
    ('PAY-18', 'Payment receipt is generated as a PDF and emailed to the vessel operator upon COMPLETED status.', 'MUST', 'Compliance'),
])

# 4.4 Settlement & Disbursement
add_heading(doc, '4.4  Settlement & Disbursement (STL)', 2)
req_table(doc, [
    ('STL-01', 'On PaymentCompleted event, the settlement engine calculates split allocations using the active SplitRule for the payment\'s FeeSchedule.', 'MUST', 'Core domain'),
    ('STL-02', 'Settlement beneficiaries for v1.0: NIMASA, Port Authority, Blue Wave, Liquid Payments. Each is a named beneficiary in the SplitRule.', 'MUST', 'Business rule'),
    ('STL-03', 'SplitRule allocations are expressed in basis points. All basis points across all beneficiaries for a given rule must sum to exactly 10,000 (= 100%). Invalid rules are rejected at save time.', 'MUST', 'Domain invariant'),
    ('STL-04', 'Each beneficiary has a Wallet. Settlement credits each beneficiary\'s Wallet with their allocated share. Ledger entries are append-only.', 'MUST', 'Ledger integrity'),
    ('STL-05', 'Collected funds are held in an escrow wallet until a SettlementBatch is triggered. Escrow is never credited below zero.', 'MUST', 'Financial safety'),
    ('STL-06', 'SettlementBatch is triggered on a configurable schedule (daily or weekly, configurable by ADMIN without code deployment).', 'MUST', 'Operations'),
    ('STL-07', 'NGN payouts are disbursed via Paystack Transfer API or Flutterwave Transfer API to registered Nigerian bank accounts.', 'MUST', 'NGN payout'),
    ('STL-08', 'USD payouts to international beneficiaries are disbursed via Wise Business API (primary). SWIFT wire is a fallback for Wise-unsupported corridors.', 'MUST', 'USD payout'),
    ('STL-09', 'The Wise Business API integration must support: balance check, recipient creation, transfer initiation, transfer status polling, and webhook confirmation.', 'MUST', 'Wise integration'),
    ('STL-10', 'Future: USD payout via stablecoin (USDT/USDC) rails. Architecture must support a pluggable payout adapter interface, same pattern as payment gateway.', 'MUST', 'Architecture'),
    ('STL-11', 'Payout failure is recorded on the SettlementBatch with error detail. Failed payouts trigger an admin alert and must be retried manually or via an admin-initiated retry action.', 'MUST', 'Resilience'),
    ('STL-12', 'Wallet available balance may never go below zero. Any operation that would result in a negative balance is rejected.', 'MUST', 'Domain invariant'),
    ('STL-13', 'All LedgerEntry records are append-only. No updates or deletes on ledger data. There is no updatedAt field on ledger entries.', 'MUST', 'Ledger integrity'),
    ('STL-14', 'A beneficiary\'s Finance Officer can view their wallet balance, ledger history, and payout history. They cannot view other beneficiaries\' data.', 'MUST', 'Data isolation'),
    ('STL-15', 'Platform Admin can view all wallets, all settlement batches, and all payout records across beneficiaries.', 'MUST', 'Admin capability'),
    ('STL-16', 'SplitRule may be versioned. A new version takes effect from a configured date. Historical PaymentOrders retain the SplitRule version active at their creation time.', 'SHOULD', 'Auditability'),
])

# 4.5 Reporting & Audit
add_heading(doc, '4.5  Reporting & Audit (RPT)', 2)
req_table(doc, [
    ('RPT-01', 'All domain events (PaymentOrderCreated, PaymentCompleted, PaymentFailed, WalletCredited, SettlementDispatched, etc.) are persisted to the reporting_domain_events outbox table before any downstream processing.', 'MUST', 'Event sourcing / audit'),
    ('RPT-02', 'The outbox pattern ensures events are not lost if downstream consumers (BullMQ workers) are unavailable.', 'MUST', 'Reliability'),
    ('RPT-03', 'Revenue summary reports: filterable by date range, port of call, vessel, beneficiary, currency, payment status.', 'MUST', 'Reporting'),
    ('RPT-04', 'Settlement batch reports: per batch, showing each beneficiary allocation, payout status, and amounts.', 'MUST', 'Reporting'),
    ('RPT-05', 'Vessel payment history: all PaymentOrders for a given IMO number, with status, amount, date, and receipt link.', 'MUST', 'Reporting'),
    ('RPT-06', 'Export formats: CSV (machine-readable) and PDF (human-readable) for all reports.', 'MUST', 'Compliance'),
    ('RPT-07', 'Audit log: every create/update action on any aggregate is logged with actor, timestamp, old value (if applicable), new value, and IP address.', 'MUST', 'NDPR / GDPR'),
    ('RPT-08', 'AUDITOR role can access all reports and the audit log but cannot export data without explicit ADMIN approval (export approval workflow).', 'COULD', 'Compliance+'),
    ('RPT-09', 'Admin dashboard: real-time revenue KPIs (total collected today/week/month, by currency, by beneficiary).', 'SHOULD', 'Operations'),
    ('RPT-10', 'Automated scheduled report delivery: configurable reports emailed to FINANCE officers on a schedule.', 'COULD', 'Nice-to-have'),
])

# 4.6 Notifications
add_heading(doc, '4.6  Notifications (NTF)', 2)
req_table(doc, [
    ('NTF-01', 'Email confirmation sent to vessel operator email on PaymentCompleted, including: amount paid, currency, Payment Ref ID, transaction reference, and PDF receipt attachment.', 'MUST', 'Core UX'),
    ('NTF-02', 'Email alert to platform ADMIN on PaymentFailed (gateway failure, not operator cancellation).', 'MUST', 'Operations'),
    ('NTF-03', 'Email alert to platform ADMIN on SettlementFailed (payout failure).', 'MUST', 'Operations'),
    ('NTF-04', 'In-app notification to FINANCE officer when a SettlementBatch is completed and wallet is credited.', 'MUST', 'UX'),
    ('NTF-05', 'Webhook callback to external emission authority on PaymentCompleted (see EMS-10).', 'MUST', 'External integration'),
    ('NTF-06', 'SMS notification to vessel operator on PaymentCompleted (configurable opt-in at org level).', 'SHOULD', 'UX'),
    ('NTF-07', 'Email reminder to vessel operator if a Payment Ref ID has been looked up but not paid within 24 hours.', 'COULD', 'Conversion'),
    ('NTF-08', 'All notification templates are configurable by ADMIN without code deployment (template stored in DB).', 'SHOULD', 'Operations'),
])

# ── 5. PAYMENT & CURRENCY REQUIREMENTS ───────────────────────────────────────
add_heading(doc, '5. Payment & Currency Requirements')

add_heading(doc, '5.1  Multi-Currency Architecture', 2)
add_body(doc,
    'The platform is USD-native in billing but multi-currency in collection. The following '
    'table summarises the currency and gateway matrix for v1.0 and planned future phases.')
currency_matrix = [
    ('USD',  'Flutterwave (primary), Paystack USD (secondary)',   'Wise Business API (primary)',  'v1.0',    'USD card / bank transfer'),
    ('NGN',  'Paystack (primary), Flutterwave (secondary)',       'Paystack Transfer API',        'v1.0',    'Card, bank transfer, USSD'),
    ('USDT', 'Crypto gateway (TBD — e.g., Coinbase Commerce)',    'On-chain transfer / off-ramp', 'Future',  'ERC-20 / TRC-20 stablecoin'),
    ('USDC', 'Crypto gateway (TBD)',                             'Circle API / on-chain',         'Future',  'ERC-20 stablecoin'),
    ('GBP',  'Flutterwave (if demand exists)',                   'Wise Business API',             'Future',  'UK-based operators'),
]
make_table(doc,
    ['Currency', 'Collection Gateway(s)', 'Payout Method', 'Phase', 'Notes'],
    currency_matrix, col_widths=[0.6, 2.0, 1.8, 0.7, 2.0], compact=True)

add_heading(doc, '5.2  FX Rate Management', 2)
fx_reqs = [
    ('FX-01', 'The FX rate (USD/NGN) is fetched from a configurable provider (e.g., ExchangeRate-API, Open Exchange Rates) at PaymentOrder creation time.', 'MUST'),
    ('FX-02', 'The fetched rate is stored on the PaymentOrder and is immutable for the lifetime of that order. Rate fluctuations after creation do not affect an in-progress payment.', 'MUST'),
    ('FX-03', 'FX provider is configurable by ADMIN. A fallback provider must be configured.', 'SHOULD'),
    ('FX-04', 'If FX rate fetch fails, NGN payment option is disabled for that session. Only USD payment is offered until rate is recoverable.', 'MUST'),
    ('FX-05', 'The FX rate shown to the operator includes a configurable spread margin (e.g., +0.5%) to offset FX risk. Margin is configurable by ADMIN.', 'SHOULD'),
]
make_table(doc, ['ID', 'Requirement', 'Priority'], fx_reqs, col_widths=[0.6, 5.8, 0.7])

add_heading(doc, '5.3  Gateway Adapter Pattern', 2)
add_body(doc,
    'All payment gateway integrations MUST conform to a common PaymentGateway interface. '
    'This enables: (a) swapping or adding gateways without modifying core payment logic, '
    '(b) routing to different gateways per currency/configuration, and (c) a future '
    'own-gateway implementation as a drop-in adapter.')
add_body(doc, 'Required interface methods:')
for method in [
    'initiatePayment(order: PaymentOrder): GatewayResponse',
    'verifyWebhook(payload: unknown, signature: string): boolean',
    'handleWebhookEvent(event: GatewayEvent): DomainEvent',
    'refundPayment(order: PaymentOrder, reason: string): RefundResponse',
    'getPaymentStatus(gatewayRef: string): PaymentStatus',
]:
    add_body(doc, method, bullet=True)

add_heading(doc, '5.4  USD Payout — Wise Business API', 2)
add_body(doc,
    'Wise Business API is selected as the primary USD payout mechanism over SWIFT wires '
    'due to lower fees, faster settlement (1–2 days vs 3–5 days), real-time rate locking, '
    'and developer-friendly API. The integration must support:')
for item in [
    'Profile verification and balance management',
    'Recipient account creation (bank account details for USD payouts to US/EU/UK accounts)',
    'Quote creation (lock exchange rate for payout amount)',
    'Transfer creation and confirmation',
    'Transfer status polling and webhook receipt for settlement confirmation',
    'Multi-currency accounts (USD, GBP, EUR) for beneficiary payouts',
]:
    add_body(doc, item, bullet=True)
add_note(doc,
    'For NIMASA and Port Authority (Nigerian government entities), NGN payouts via Paystack/Flutterwave '
    'to Nigerian bank accounts are the primary mechanism. USD Wise payouts are primarily for '
    'Blue Wave and Liquid Payments international disbursements.')

# ── 6. NON-FUNCTIONAL REQUIREMENTS ───────────────────────────────────────────
add_heading(doc, '6. Non-Functional Requirements')

add_heading(doc, '6.1  Performance', 2)
perf = [
    ('NFR-PERF-01', 'API p95 response time ≤ 500ms for all endpoints excluding external API calls.', 'MUST'),
    ('NFR-PERF-02', 'Payment webhook processing (BullMQ job) completed within 5 seconds of receipt.', 'MUST'),
    ('NFR-PERF-03', 'External API lookup (Payment Ref ID resolution) timeout: 10 seconds. Circuit breaker opens after 3 consecutive failures.', 'MUST'),
    ('NFR-PERF-04', 'Dashboard reports for up to 12 months of data must render within 3 seconds.', 'SHOULD'),
    ('NFR-PERF-05', 'System must sustain 200 concurrent payment sessions without degradation.', 'SHOULD'),
]
make_table(doc, ['ID', 'Requirement', 'Priority'], perf, col_widths=[1.3, 5.3, 0.7])

add_heading(doc, '6.2  Security', 2)
sec = [
    ('NFR-SEC-01', 'All data in transit encrypted via TLS 1.3 minimum. TLS 1.1 and 1.2 disabled.', 'MUST'),
    ('NFR-SEC-02', 'All PII and payment data at rest encrypted using AES-256.', 'MUST'),
    ('NFR-SEC-03', 'JWT RS256 asymmetric signing. Private key stored in environment secrets — never in source code or DB.', 'MUST'),
    ('NFR-SEC-04', 'TOTP 2FA mandatory for ADMIN role. FINANCE and OPERATOR roles may optionally enable 2FA.', 'MUST'),
    ('NFR-SEC-05', 'No card data stored on platform servers. PCI-DSS SAQ-A scope (gateway-hosted card capture).', 'MUST'),
    ('NFR-SEC-06', 'Webhook HMAC signature verification for all incoming gateway webhooks.', 'MUST'),
    ('NFR-SEC-07', 'Rate limiting: 100 req/min per authenticated user; 20 req/min for unauthenticated endpoints.', 'MUST'),
    ('NFR-SEC-08', 'OWASP Top 10 mitigations applied: input validation, parameterised queries, XSS prevention, CSRF tokens, secure headers (CSP, HSTS).', 'MUST'),
    ('NFR-SEC-09', 'Dependency vulnerability scanning in CI (e.g., npm audit, Snyk) on every pull request.', 'MUST'),
    ('NFR-SEC-10', 'Penetration test by a third party required before production launch (Phase 4 gate).', 'MUST'),
]
make_table(doc, ['ID', 'Requirement', 'Priority'], sec, col_widths=[1.3, 5.3, 0.7])

add_heading(doc, '6.3  Availability & Reliability', 2)
avail = [
    ('NFR-AV-01', 'Target uptime: 99.9% per calendar month (≤ 43.8 minutes downtime/month).', 'MUST'),
    ('NFR-AV-02', 'Planned maintenance windows must be communicated 48 hours in advance and scheduled outside Nigerian business hours (06:00–22:00 WAT).', 'SHOULD'),
    ('NFR-AV-03', 'BullMQ jobs have a minimum of 3 retry attempts with exponential backoff before marking failed.', 'MUST'),
    ('NFR-AV-04', 'Database: automated daily backups with 30-day retention. Point-in-time recovery capability required.', 'MUST'),
    ('NFR-AV-05', 'Circuit breaker pattern on all external API integrations (emission authority, payment gateways, payout providers, FX providers).', 'MUST'),
]
make_table(doc, ['ID', 'Requirement', 'Priority'], avail, col_widths=[1.3, 5.3, 0.7])

add_heading(doc, '6.4  Scalability', 2)
scale = [
    ('NFR-SC-01', 'The API layer must be horizontally scalable (stateless; session state in Redis, not in-process memory).', 'MUST'),
    ('NFR-SC-02', 'BullMQ workers must be independently scalable from the API layer.', 'MUST'),
    ('NFR-SC-03', 'Database connection pooling (PgBouncer or Prisma connection pool) required in production.', 'MUST'),
    ('NFR-SC-04', 'Architecture must support processing ≥ 1,000 payment webhook events per hour without queue backlog.', 'SHOULD'),
]
make_table(doc, ['ID', 'Requirement', 'Priority'], scale, col_widths=[1.3, 5.3, 0.7])

add_heading(doc, '6.5  Compliance', 2)
comp = [
    ('NFR-CP-01', 'NDPR 2019 compliance: lawful basis for processing, data subject rights (access, deletion, portability), DPO designation, breach notification within 72 hours.', 'MUST'),
    ('NFR-CP-02', 'GDPR compliance for EU-based vessel operators.', 'MUST'),
    ('NFR-CP-03', 'PCI-DSS SAQ-A: no card data on platform servers; quarterly ASV scans on in-scope components.', 'MUST'),
    ('NFR-CP-04', 'All financial records retained for a minimum of 7 years (FIRS requirement).', 'MUST'),
    ('NFR-CP-05', 'MARPOL Annex VI emission calculation methodology documented and auditable.', 'MUST'),
    ('NFR-CP-06', 'DPIA completed before any PII data is collected in production (Phase 4 gate).', 'MUST'),
]
make_table(doc, ['ID', 'Requirement', 'Priority'], comp, col_widths=[1.3, 5.3, 0.7])

# ── 7. DATA REQUIREMENTS ─────────────────────────────────────────────────────
add_heading(doc, '7. Data Requirements')

add_heading(doc, '7.1  Data Principles', 2)
for item in [
    'All monetary amounts stored as integer minor units (cents for USD, kobo for NGN) with a Currency enum field. No floating-point money.',
    'No cross-context foreign key constraints in the database. Cross-context references use plain string IDs.',
    'Table names are prefixed by context: identity_, emissions_, payments_, settlement_, reporting_, notifications_.',
    'Ledger entries (settlement_ledger_entries) have no updatedAt field. Append-only.',
    'SplitRule basis points must sum to 10,000; enforced at application layer and documented as a DB check constraint.',
    'All timestamps stored in UTC. Timezone conversion is a presentation-layer concern.',
    'PII fields (email, phone, full name, address) encrypted at rest using AES-256.',
]:
    add_body(doc, item, bullet=True)

add_heading(doc, '7.2  Data Retention', 2)
retention = [
    ('Payment records',          '7 years',      'FIRS / CBN regulatory requirement'),
    ('Settlement & ledger',      '7 years',      'Financial audit requirement'),
    ('Emission assessment data', '7 years',      'NIMASA / MARPOL record-keeping'),
    ('Audit logs',               '5 years',      'NDPR / GDPR'),
    ('Domain events (outbox)',   '2 years active + archive', 'Operational + compliance'),
    ('Notification logs',        '1 year',       'Operational'),
    ('Session tokens (Redis)',   '7 days',        'Security'),
    ('Idempotency keys (Redis)', '24 hours',      'Operational'),
]
make_table(doc, ['Data Category', 'Retention Period', 'Rationale'], retention, col_widths=[2.0, 1.8, 3.2])

add_heading(doc, '7.3  Data Sovereignty', 2)
add_body(doc,
    'All primary data must be stored in servers physically located in Nigeria or, at minimum, '
    'within Africa, to comply with NDPR data localisation requirements. The Railway-managed '
    'PostgreSQL instance must be provisioned in a Nigerian or West African region. If no such '
    'region is available on Railway, an alternative provider (e.g., AWS Lagos ap-southeast-1 '
    'equivalent, or Render) must be evaluated.')

# ── 8. INTERFACE REQUIREMENTS ─────────────────────────────────────────────────
add_heading(doc, '8. Interface Requirements')

add_heading(doc, '8.1  User Interface', 2)
for item in [
    'The web application is responsive. All payment flows must be fully functional on mobile browsers (iOS Safari, Android Chrome).',
    'UI is built with Next.js 14 (App Router) + shadcn/ui + Tailwind CSS.',
    'All monetary amounts displayed with currency symbol, formatted to 2 decimal places.',
    'Multi-currency payment selection presents the USD amount and the NGN equivalent (with "rate locked at" disclosure) side-by-side.',
    'All forms include client-side and server-side validation. Error messages are human-readable and actionable.',
    'Loading states are shown for all async operations (API lookups, payment processing).',
]:
    add_body(doc, item, bullet=True)

add_heading(doc, '8.2  External API Interfaces', 2)
ext_apis = [
    ('Emission Authority API', 'REST (assumed) or SOAP', 'GET assessment data by Payment Ref ID; POST payment confirmation', 'TLS 1.3; API key or mutual TLS'),
    ('Paystack',               'REST',                    'Payment initiation, webhook events, transfer (NGN payout)',         'HTTPS; HMAC-SHA512 webhook verification'),
    ('Flutterwave',            'REST',                    'Payment initiation (USD + NGN), webhook events, transfer',          'HTTPS; SHA-256 webhook hash verification'),
    ('Wise Business API',      'REST',                    'Quote, transfer creation, recipient management, webhook',           'OAuth 2.0; HTTPS'),
    ('FX Rate Provider',       'REST',                    'USD/NGN spot rate fetch',                                           'HTTPS; API key'),
    ('Email Provider',         'REST (SendGrid/Mailgun)', 'Transactional email delivery',                                      'HTTPS; API key'),
    ('SMS Provider',           'REST (Termii)',            'SMS notification delivery (optional)',                              'HTTPS; API key'),
]
make_table(doc,
    ['System', 'Protocol', 'Operations', 'Auth / Security'],
    ext_apis, col_widths=[1.3, 1.0, 2.7, 2.0], compact=True)

add_heading(doc, '8.3  Internal API', 2)
add_body(doc,
    'The platform exposes a REST API documented in OpenAPI 3.x (YAML). Key API surface areas:')
api_groups = [
    ('Auth',        '/auth/register, /auth/login, /auth/refresh, /auth/logout, /auth/2fa/*'),
    ('Identity',    '/organisations, /organisations/{id}/vessels, /users, /users/{id}'),
    ('Assessment',  '/assessments/lookup (POST with paymentRefId), /assessments/{id}'),
    ('Payments',    '/payment-orders, /payment-orders/{id}, /payment-orders/{id}/initiate, /webhooks/{gateway}'),
    ('Settlement',  '/wallets, /wallets/{id}/ledger, /settlement-batches, /payouts'),
    ('Reports',     '/reports/revenue, /reports/settlement, /reports/vessels/{imoNumber}'),
    ('Admin',       '/admin/organisations, /admin/fee-schedules, /admin/split-rules, /admin/gateways, /admin/templates'),
]
make_table(doc, ['API Group', 'Key Endpoints'], api_groups, col_widths=[1.4, 5.6])

# ── 9. CONSTRAINTS ────────────────────────────────────────────────────────────
add_heading(doc, '9. System Constraints')
constraints = [
    ('Technical',     'TypeScript (strict) throughout — no plain JavaScript in any deliverable.'),
    ('Technical',     'Monorepo structure: Turborepo. Apps: apps/web (Next.js), apps/api (Express). Shared packages: packages/ui, packages/types, packages/config.'),
    ('Technical',     'ORM: Prisma 5.x only. No raw SQL except for complex reporting queries (documented exceptions).'),
    ('Technical',     'No floating-point arithmetic on monetary values anywhere in the codebase.'),
    ('Regulatory',    'Card data must not pass through platform servers. PCI-DSS SAQ-A only.'),
    ('Regulatory',    'PII must not appear in application logs (Pino) or error tracking. Use redaction middleware.'),
    ('Regulatory',    'Production deployment requires DPIA sign-off and pen-test report.'),
    ('Business',      'The platform is not the source of truth for emission assessments. The external maritime authority system is. Disputes on assessment amounts are handled externally.'),
    ('Business',      'The platform cannot override or modify the base emission fee from the external system (admin override in EMS-13 requires full audit trail and is for exceptional cases only).'),
    ('Business',      'Blue Wave is the operating entity. NIMASA and Port Authority are beneficiaries, not operators of the platform.'),
]
make_table(doc, ['Constraint Type', 'Description'], constraints, col_widths=[1.3, 5.7])

# ── 10. OPEN ITEMS & FUTURE PHASES ───────────────────────────────────────────
add_heading(doc, '10. Open Items & Future Phases')

add_heading(doc, '10.1  Open Items Requiring Decision Before Build', 2)
open_items = [
    ('OI-01', 'External Emission Authority API — MOCKED',
     'Full mock built at /api/assessment/[refId]. Payload schema defined (14 fields inc. feeAmountUSD, co2eTons, assessmentStatus). 4 test vessels registered. Production integration requires real API credentials + endpoint from NIMASA.',
     'IN PROGRESS — mock live, prod creds needed'),
    ('OI-02', 'FX Rate Provider — RESOLVED',
     'ExchangeRate-API selected. Spot USD/NGN rate mocked at ₦1,637.20. Platform applies 1.5% spread → effective rate ₦1,661.72. Rate locked per PaymentOrder for 30 minutes.',
     'RESOLVED'),
    ('OI-03', 'USD Payout Provider — RESOLVED',
     'Grey.co Business API selected. Provides USD virtual accounts and programmatic cross-border transfers for Blue Wave Maritime and Liquid Payments. NGN beneficiaries (NIMASA, Port Authority) paid via Paystack Transfer API.',
     'RESOLVED'),
    ('OI-04', 'SplitRule Configuration — RESOLVED',
     'Equal split confirmed: 25% each (2,500 basis points) to NIMASA, Lagos Port Authority, Blue Wave Maritime, Liquid Payments. Total = 10,000 bp = 100%. Four parties, no reserve fund.',
     'RESOLVED'),
    ('OI-05', 'Platform Service Charge — RESOLVED',
     '1.2% of base emission fee (percentage, not flat). Example: $124,730 base × 1.2% = $1,496.76 platform charge. Total payable = $126,226.76.',
     'RESOLVED'),
    ('OI-06', 'Email & SMS Providers — RESOLVED',
     'Email: Resend (resend.com) — from noreply@liquidemissions.ng, 7 transactional templates. SMS: Termii — sender ID LQDEMIT, 3 SMS templates, optional per-organisation.',
     'RESOLVED'),
]
make_table(doc,
    ['ID', 'Item', 'Detail', 'Impact'],
    open_items, col_widths=[0.55, 1.7, 4.2, 0.85 ], compact=True)

add_heading(doc, '10.2  Future Phases (Post v1.0)', 2)
future = [
    ('F-01', 'Crypto Collection',    'USDT/USDC payment via Coinbase Commerce or Bitpay. Stablecoin payout to beneficiaries.'),
    ('F-02', 'Own Payment Gateway',  'Proprietary gateway replacing Paystack/Flutterwave. Pluggable adapter pattern (PAY-15) enables this with minimal core changes.'),
    ('F-03', 'AIS/SEMSW Integration','Direct ingestion of vessel emission data from physical sensors. Eliminates reliance on external authority API for data.'),
    ('F-04', 'Native Mobile Apps',   'iOS and Android apps for vessel operators.'),
    ('F-05', 'Multi-Jurisdiction',   'Extend to other West African port jurisdictions beyond Nigeria.'),
    ('F-06', 'SWIFT Alternative',    'If Wise is unavailable: evaluate Grey.co, Leatherback, or Flutterwave\'s international USD payout for beneficiaries requiring USD outside Wise corridors.'),
    ('F-07', 'Instalment Payments',  'For large assessments, allow phased payment plans (requires escrow logic extension).'),
    ('F-08', 'Advanced Analytics',   'BI dashboard with trend analysis, port comparisons, and predictive revenue modelling.'),
]
make_table(doc, ['ID', 'Feature', 'Description'], future, col_widths=[0.5, 1.8, 4.7])

# ── APPENDIX ──────────────────────────────────────────────────────────────────
add_heading(doc, 'Appendix A — Payment State Machine')
add_body(doc, 'Valid state transitions for PaymentOrder:')
states = [
    ('PENDING',             'PROCESSING',  'Payment initiation acknowledged by gateway'),
    ('PROCESSING',          'COMPLETED',   'Gateway webhook confirms successful charge'),
    ('PROCESSING',          'FAILED',      'Gateway webhook confirms failed or declined charge'),
    ('COMPLETED',           'REFUNDED',    'Refund processed via new reversal PaymentOrder (original remains COMPLETED + immutable)'),
    ('FAILED',              'PENDING',     'Operator retries payment (new PaymentOrder created referencing same Payment Ref ID — idempotency check passes because prior order is FAILED, not COMPLETED)'),
    ('PENDING/PROCESSING',  'CANCELLED',   'Operator cancels before payment submitted to gateway (PROCESSING → CANCELLED only within 30s grace window via gateway void)'),
]
make_table(doc, ['From State', 'To State', 'Trigger / Condition'], states, col_widths=[1.4, 1.2, 4.4])

add_heading(doc, 'Appendix B — Domain Events Reference')
events = [
    ('PaymentOrderCreated',    'payments',   'New PaymentOrder created from assessment lookup'),
    ('PaymentInitiated',       'payments',   'Gateway payment session started'),
    ('PaymentCompleted',       'payments',   'Gateway confirms successful payment'),
    ('PaymentFailed',          'payments',   'Gateway confirms failed payment'),
    ('PaymentRefunded',        'payments',   'Reversal PaymentOrder completed'),
    ('WalletCredited',         'settlement', 'Beneficiary wallet credited after settlement split'),
    ('EscrowFunded',           'settlement', 'Escrow wallet funded on PaymentCompleted'),
    ('SettlementBatchCreated', 'settlement', 'New batch triggered (manual or scheduled)'),
    ('SettlementDispatched',   'settlement', 'Payout API call made for beneficiary'),
    ('SettlementCompleted',    'settlement', 'Payout confirmed by provider webhook'),
    ('SettlementFailed',       'settlement', 'Payout failed — admin alert triggered'),
    ('AssessmentRetrieved',    'emissions',  'External API returned assessment data'),
    ('AssessmentIncomplete',   'emissions',  'External API returned insufficient data for fee calculation'),
    ('UserRegistered',         'identity',   'New user account created'),
    ('OrganisationCreated',    'identity',   'New organisation onboarded'),
    ('OrganisationSuspended',  'identity',   'Organisation suspended by admin'),
]
make_table(doc, ['Event Name', 'Context', 'Description'], events, col_widths=[2.2, 1.1, 3.7])

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = r'c:\ReveenueCollect\docs\Liquid_Emissions_SRS.docx'
doc.save(out)
print(f'Saved: {out}')
