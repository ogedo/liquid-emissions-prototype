"""
Generate MetronBase Labs Product Development Playbook
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x3E)   # dark navy
TEAL   = RGBColor(0x00, 0x8B, 0x8B)   # teal accent
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xF8)   # light blue-grey for alt rows
DARK_T = RGBColor(0x1A, 0x2E, 0x4A)   # header rows in tables

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), val.get('val', 'single'))
        border.set(qn('w:sz'), str(val.get('sz', 4)))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def para_font(para, size=10, bold=False, italic=False,
              color: RGBColor = None, align=None):
    if align:
        para.alignment = align
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

def add_heading(doc, text, level=1):
    """Add a styled heading with navy/teal accent."""
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = NAVY
        # bottom border accent
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '008B8B')
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = TEAL
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = DARK_T
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

def add_body(doc, text, italic=False, bullet=False):
    if bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def table_header_row(table, headers, col_widths=None):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, DARK_T)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(hdr)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            cell.width = Inches(col_widths[i])

def add_table_row(table, values, alt=False, col_widths=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if alt:
            set_cell_bg(cell, LIGHT)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def make_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_header_row(tbl, headers, col_widths)
    for idx, row in enumerate(rows):
        add_table_row(tbl, row, alt=(idx % 2 == 1), col_widths=col_widths)
    doc.add_paragraph()
    return tbl

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        run = p.add_run('MetronBase Labs  |  Product Development Playbook v1.0  |  Confidential — Internal Use Only')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT BUILD
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

add_footer(doc)

# ── COVER PAGE ───────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)

# navy block — simulate with heavily styled paragraph
def cover_block(doc):
    # Agency name
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run('MetronBase Labs')
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Product Development Playbook')
    r2.font.size = Pt(22)
    r2.font.bold = True
    r2.font.color.rgb = TEAL

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('Version 1.0')
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p3.paragraph_format.space_after = Pt(8)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run('"From Discovery to Deployment — The MetronBase Way"')
    r4.font.size = Pt(13)
    r4.font.italic = True
    r4.font.color.rgb = DARK_T
    p4.paragraph_format.space_after = Pt(40)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run('Confidential — Internal Use Only')
    r5.font.size = Pt(9)
    r5.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

cover_block(doc)
doc.add_page_break()

# ── SECTION 1: About This Playbook ──────────────────────────────────────────
add_heading(doc, '1. About This Playbook')
add_body(doc, 'This playbook is MetronBase Labs\' authoritative reference for how we design, build, '
         'and ship digital products. It defines our document framework, phase-by-phase process, '
         'quality gates, and tooling standards.')

add_heading(doc, '1.1 Purpose & Scope', 2)
add_body(doc, 'The Playbook governs every client engagement from initial discovery through to post-launch '
         'handoff. It ensures consistency, auditability, and a high bar of engineering quality across all '
         'projects regardless of size or domain.')

add_heading(doc, '1.2 Audience', 2)
for role in [
    'Project Managers — phase tracking, gate sign-off, document ownership',
    'Engineers (Frontend, Backend, DevOps) — coding standards, architecture expectations',
    'Designers — UI/UX specification requirements and deliverable format',
    'QA Engineers — test plan ownership and quality gate criteria',
    'Clients — transparency into our process and what to expect at each phase',
]:
    add_body(doc, role, bullet=True)

add_heading(doc, '1.3 How to Use This Document', 2)
add_body(doc, 'Navigate by phase (Section 4) to understand what is required at each stage. '
         'Use Section 3 as the definitive list of all documents. Use Section 5 to verify gate '
         'criteria before advancing. For fast-track MVPs, skip to Section 8.')

# ── SECTION 2: Core Principles ──────────────────────────────────────────────
add_heading(doc, '2. Core Principles')

principles = [
    ('Domain-First Thinking (DDD)',
     'We model the business domain before writing a single line of code. Bounded contexts, '
     'aggregates, and domain events are defined in collaboration with stakeholders before architecture decisions are made.'),
    ('Documentation as Code',
     'Docs live in the repository alongside the code they describe. They are versioned, reviewed '
     'in PRs, and kept up to date as a first-class engineering artifact.'),
    ('Just-in-Time Documents',
     'We produce documentation per phase, not all upfront. A Tier 6 runbook is not written during '
     'discovery. This keeps docs relevant, accurate, and unbloated.'),
    ('Security & Compliance by Default',
     'Security architecture, data privacy impact assessments, and compliance checklists are not '
     'afterthoughts — they are gated requirements before any production deployment.'),
    ('Iterative Delivery',
     'We ship in increments. Each sprint produces working software and updated documentation. '
     'No big-bang releases.'),
    ('Single Source of Truth',
     'One canonical document per concern. No parallel versions, no stale copies. The repository '
     'is the source of truth.'),
]

for title, body in principles:
    add_heading(doc, title, 3)
    add_body(doc, body)

# ── SECTION 3: Document Framework ───────────────────────────────────────────
add_heading(doc, '3. The Document Framework (7 Tiers)')
add_body(doc, 'All MetronBase Labs projects follow a seven-tier document hierarchy. Each tier maps to '
         'a set of artifacts with defined owners, phases, and priority levels.')

tier_headers = ['Document Name', 'Purpose', 'Owner', 'Phase Required', 'Priority']
tier_widths  = [1.5, 2.4, 0.9, 1.1, 1.1]

tiers = [
    {
        'name': 'Tier 1 — Strategy & Discovery',
        'rows': [
            ['Product Vision Doc',   'Defines the "why" — mission, goals, success metrics',          'PM',        'Phase 1', 'MVP'],
            ['Business Requirements\nDocument (BRD)', 'Captures business objectives, constraints, stakeholders', 'PM', 'Phase 1', 'MVP'],
            ['Market & Competitive\nAnalysis',        'Validates market fit and differentiators',                'PM',        'Phase 1', 'Nice-to-Have'],
        ],
    },
    {
        'name': 'Tier 2 — Requirements',
        'rows': [
            ['Product Requirements\nDocument (PRD)', 'Features, user flows, functional requirements',            'PM',        'Phase 2', 'MVP'],
            ['Software Requirements\nSpec (SRS)',    'Technical requirements, system constraints, use cases',    'Tech Lead', 'Phase 2', 'MVP'],
            ['User Stories / Backlog',               'Granular, sprint-ready work items',                        'PM',        'Phase 2', 'MVP'],
            ['API Specification\n(OpenAPI 3.x)',     'Machine-readable contract for all API endpoints',          'Tech Lead', 'Phase 2', 'MVP'],
        ],
    },
    {
        'name': 'Tier 3 — Architecture & Design',
        'rows': [
            ['Domain Model',                 'DDD strategic + tactical map — contexts, aggregates, events', 'Tech Lead', 'Phase 2', 'MVP'],
            ['System Architecture Doc\n(SAD)', 'Components, integrations, deployment topology',             'Tech Lead', 'Phase 2', 'MVP'],
            ['Database Schema',              'Prisma schema, ER diagram, migration strategy',               'Tech Lead', 'Phase 2', 'MVP'],
            ['UI/UX Specification',          'Wireframes, design system, component map',                    'Designer',  'Phase 2', 'MVP'],
        ],
    },
    {
        'name': 'Tier 4 — Security & Compliance',
        'rows': [
            ['Security Architecture Doc',   'Threat model, auth design, encryption standards',            'Tech Lead', 'Phase 2', 'MVP'],
            ['Data Privacy Impact\nAssessment (DPIA)', 'GDPR/NDPR data flows, risk register, DPO sign-off', 'PM',        'Phase 4', 'MVP'],
            ['PCI-DSS / Payment\nCompliance Checklist', 'Compliance controls for payment card data',        'Tech Lead', 'Phase 4', 'Launch'],
            ['Pen-Test Report',              'Third-party penetration test findings and remediations',     'DevOps',    'Phase 4', 'Launch'],
        ],
    },
    {
        'name': 'Tier 5 — Development Operations',
        'rows': [
            ['Coding Standards Guide',      'TypeScript style, naming conventions, linting rules',        'Tech Lead', 'Phase 3', 'MVP'],
            ['Git Strategy Doc',            'Branching model, commit conventions, PR workflow',            'Tech Lead', 'Phase 3', 'MVP'],
            ['Test Plan',                   'Unit, integration, E2E coverage targets and tooling',         'QA',        'Phase 3', 'MVP'],
            ['Environment Setup Guide',     'Local dev, staging, and CI environment configuration',        'DevOps',    'Phase 3', 'MVP'],
        ],
    },
    {
        'name': 'Tier 6 — Deployment & Operations',
        'rows': [
            ['Deployment Runbook',          'Step-by-step production deployment procedure',               'DevOps',    'Phase 5', 'Launch'],
            ['Infrastructure Diagram',      'Cloud topology, network layout, service dependencies',        'DevOps',    'Phase 5', 'Launch'],
            ['Monitoring & Alerting Plan',  'Metrics, dashboards, alert thresholds, on-call rota',         'DevOps',    'Phase 5', 'Launch'],
            ['Incident Response Playbook',  'Severity definitions, escalation matrix, post-mortem template', 'DevOps', 'Phase 5', 'Launch'],
        ],
    },
    {
        'name': 'Tier 7 — Handoff & Maintenance',
        'rows': [
            ['API Reference Docs',          'End-user API documentation (e.g., Redoc / Stoplight)',        'Tech Lead', 'Phase 5', 'Launch'],
            ['Developer Onboarding Guide',  'Repo walkthrough, setup steps, architectural overview',       'Tech Lead', 'Phase 5', 'Launch'],
            ['User Manual',                 'End-user guide for client-facing features',                   'PM',        'Phase 5', 'Nice-to-Have'],
            ['Operations Manual',           'Day-to-day ops, admin tasks, backup & restore procedures',    'DevOps',    'Phase 5', 'Nice-to-Have'],
        ],
    },
]

for tier in tiers:
    add_heading(doc, tier['name'], 2)
    make_table(doc, tier_headers, tier['rows'], col_widths=tier_widths)

# ── SECTION 4: Phase-by-Phase Process ────────────────────────────────────────
add_heading(doc, '4. Phase-by-Phase Process')
add_body(doc, 'All projects advance through five phases. Each phase has a clear goal, required inputs, '
         'produced outputs, and a quality gate that must pass before the next phase begins.')

phases = [
    {
        'name': 'Phase 1 — Discovery & Strategy',
        'goal': 'Align on the "why". Understand the domain, users, constraints, and market before any design or code.',
        'inputs': ['Client brief / RFP', 'Stakeholder interviews', 'Competitive research'],
        'outputs': ['Product Vision Doc', 'Business Requirements Document (BRD)', 'Initial Domain Glossary', 'Market Analysis (if required)'],
        'gate': 'BRD signed off by client. Domain glossary agreed. Project kick-off meeting completed.',
    },
    {
        'name': 'Phase 2 — Requirements & Design',
        'goal': 'Define exactly what will be built and how it will be structured — functionally and architecturally.',
        'inputs': ['Approved BRD', 'Domain expert access', 'Design system assets'],
        'outputs': ['PRD', 'SRS', 'User Stories', 'Domain Model', 'SAD', 'DB Schema', 'API Spec', 'UI/UX Spec', 'Security Architecture Doc'],
        'gate': 'PRD + SRS approved. Domain model reviewed by Tech Lead. DB schema finalised. UI/UX prototype signed off by client.',
    },
    {
        'name': 'Phase 3 — Build (Iterative Sprints)',
        'goal': 'Deliver working, tested software incrementally. Each sprint produces releasable increments.',
        'inputs': ['Approved PRD', 'SRS', 'Domain Model', 'API Spec', 'DB Schema'],
        'outputs': ['Working software (per sprint)', 'Unit & integration tests', 'Updated API spec', 'Coding Standards Guide', 'Git Strategy Doc', 'Test Plan', 'Env Setup Guide'],
        'gate': 'All sprint goals met. Test coverage targets hit. No critical open bugs. API spec matches implementation.',
    },
    {
        'name': 'Phase 4 — QA & Security',
        'goal': 'Validate quality, security, and compliance before any production exposure.',
        'inputs': ['Built application', 'Test Plan', 'Security Architecture Doc'],
        'outputs': ['QA sign-off report', 'DPIA', 'PCI-DSS Checklist (if applicable)', 'Pen-Test Report', 'Security remediation log'],
        'gate': 'All critical/high-severity defects resolved. DPIA approved by DPO. Pen-test remediation complete. Compliance checklist signed.',
    },
    {
        'name': 'Phase 5 — Launch & Handoff',
        'goal': 'Ship to production and transfer ownership with full documentation.',
        'inputs': ['QA sign-off', 'Approved compliance docs', 'Infrastructure provisioned'],
        'outputs': ['Deployment Runbook', 'Infrastructure Diagram', 'Monitoring & Alerting Plan', 'Incident Response Playbook', 'API Docs', 'Developer Onboarding Guide', 'User Manual', 'Ops Manual'],
        'gate': 'Production deployment verified. Monitoring active. Runbook walkthrough completed with client ops team. All Tier 7 docs delivered.',
    },
]

phase_headers = ['Attribute', 'Detail']
for phase in phases:
    add_heading(doc, phase['name'], 2)
    rows = [
        ['Goal', phase['goal']],
        ['Documents Required In', ', '.join(phase['inputs'])],
        ['Documents Produced', ', '.join(phase['outputs'])],
        ['Quality Gate', phase['gate']],
    ]
    make_table(doc, phase_headers, rows, col_widths=[1.4, 5.6])

# ── SECTION 5: Quality Gates ──────────────────────────────────────────────────
add_heading(doc, '5. Quality Gates')
add_body(doc, 'A quality gate is a mandatory checkpoint between phases. No phase may begin until its '
         'predecessor gate has been formally passed and recorded.')

gate_headers = ['Gate Name', 'Phase Transition', 'Required Documents', 'Sign-off Required From']
gate_rows = [
    ['Strategy Gate',     'Phase 1 → Phase 2', 'BRD, Vision Doc, Domain Glossary',                                          'PM + Client'],
    ['Design Gate',       'Phase 2 → Phase 3', 'PRD, SRS, Domain Model, SAD, DB Schema, API Spec, UI/UX Spec',               'PM + Tech Lead + Client'],
    ['Build Gate',        'Phase 3 → Phase 4', 'Sprint demos, Test Plan, updated API Spec, Coding Standards adherence report','Tech Lead + QA'],
    ['Security Gate',     'Phase 4 → Phase 5', 'DPIA, Pen-Test Report, PCI Checklist (if applicable), QA sign-off',          'Tech Lead + DPO + Client'],
    ['Launch Gate',       'Phase 5 → Live',    'Runbook, Infra Diagram, Monitoring Plan, all Tier 7 docs',                   'PM + DevOps + Client'],
]
make_table(doc, gate_headers, gate_rows, col_widths=[1.2, 1.3, 2.8, 1.7])

# ── SECTION 6: Document Ownership & RACI ──────────────────────────────────────
add_heading(doc, '6. Document Ownership & RACI')
add_body(doc, 'R = Responsible (produces it)  |  A = Accountable (approves it)  |  C = Consulted  |  I = Informed')

raci_headers = ['Document', 'PM', 'Tech Lead', 'Designer', 'QA', 'DevOps', 'Client']
raci_rows = [
    ['Product Vision Doc',           'R/A', 'C',   'C',  'I',  'I',   'A'],
    ['BRD',                          'R/A', 'C',   'I',  'I',  'I',   'A'],
    ['PRD',                          'R/A', 'C',   'C',  'C',  'I',   'A'],
    ['SRS',                          'C',   'R/A', 'I',  'C',  'C',   'I'],
    ['User Stories',                 'R',   'C',   'C',  'C',  'I',   'A'],
    ['API Specification',            'I',   'R/A', 'I',  'C',  'C',   'I'],
    ['Domain Model',                 'C',   'R/A', 'I',  'I',  'I',   'C'],
    ['System Architecture Doc',      'I',   'R/A', 'I',  'C',  'C',   'I'],
    ['Database Schema',              'I',   'R/A', 'I',  'I',  'C',   'I'],
    ['UI/UX Specification',          'C',   'C',   'R/A','C',  'I',   'A'],
    ['Security Architecture Doc',    'I',   'R/A', 'I',  'C',  'C',   'I'],
    ['DPIA',                         'R/A', 'C',   'I',  'I',  'C',   'A'],
    ['PCI-DSS Checklist',            'C',   'R',   'I',  'C',  'R',   'A'],
    ['Pen-Test Report',              'I',   'A',   'I',  'I',  'R',   'I'],
    ['Coding Standards Guide',       'I',   'R/A', 'I',  'C',  'C',   'I'],
    ['Git Strategy Doc',             'I',   'R/A', 'I',  'C',  'C',   'I'],
    ['Test Plan',                    'C',   'C',   'I',  'R/A','I',   'I'],
    ['Environment Setup Guide',      'I',   'C',   'I',  'C',  'R/A', 'I'],
    ['Deployment Runbook',           'I',   'C',   'I',  'C',  'R/A', 'I'],
    ['Infrastructure Diagram',       'I',   'C',   'I',  'I',  'R/A', 'I'],
    ['Monitoring & Alerting Plan',   'I',   'C',   'I',  'C',  'R/A', 'I'],
    ['Incident Response Playbook',   'I',   'A',   'I',  'I',  'R',   'I'],
    ['API Reference Docs',           'I',   'R/A', 'I',  'C',  'I',   'I'],
    ['Developer Onboarding Guide',   'I',   'R/A', 'I',  'I',  'C',   'I'],
    ['User Manual',                  'R/A', 'C',   'C',  'I',  'I',   'A'],
    ['Operations Manual',            'I',   'C',   'I',  'I',  'R/A', 'A'],
]
make_table(doc, raci_headers, raci_rows, col_widths=[2.2, 0.55, 0.85, 0.75, 0.55, 0.75, 0.65])

# ── SECTION 7: Tools & Standards ─────────────────────────────────────────────
add_heading(doc, '7. Tools & Standards')
add_body(doc, 'All MetronBase Labs projects default to the following technology stack and tooling. '
         'Deviations require written justification in the SAD.')

tools_headers = ['Category', 'Tool / Standard', 'Notes']
tools_rows = [
    ['Monorepo',            'Turborepo',                   'All apps and shared packages in a single repo'],
    ['Language',            'TypeScript (strict mode)',    'No plain JS. Strict null checks enabled.'],
    ['Frontend',            'Next.js 14 (App Router)',     'React Server Components by default'],
    ['Backend',             'Node.js + Express',           'REST API; structured as DDD modules'],
    ['Database',            'PostgreSQL 15',               'Managed via Railway or equivalent'],
    ['ORM',                 'Prisma 5.x',                  'prisma-client-js; migrations in version control'],
    ['Queue / Cache',       'Redis + BullMQ',              'Job queues for async tasks; cache for idempotency keys'],
    ['Payments',            'Paystack / Flutterwave',      'Dual-gateway; Nigerian NGN primary'],
    ['Architecture Style',  'DDD Modular Monolith',        'Bounded contexts as modules under apps/api/src/modules/'],
    ['UI Components',       'shadcn/ui',                   'Tailwind CSS; packages/ui/ shared component lib'],
    ['Logging',             'Pino',                        'Structured JSON logs; correlation IDs on all requests'],
    ['Documentation Fmt',   'Markdown → PDF / DOCX',       'Docs live in docs/ directory of the repository'],
    ['Version Control',     'Git (Conventional Commits)',  'feat:, fix:, chore:, docs:, test: prefixes enforced'],
    ['CI/CD',               'GitHub Actions',              'Lint → test → build → deploy pipeline per environment'],
    ['API Specification',   'OpenAPI 3.x (YAML)',          'Generated Redoc site for client-facing reference'],
    ['Auth',                'JWT RS256 + TOTP 2FA',        '2FA mandatory for admin roles'],
    ['Secrets Management',  'Environment Variables + Vault','No secrets in source control'],
    ['Compliance',          'NDPR + GDPR',                 'DPIA mandatory for all data-handling features'],
]
make_table(doc, tools_headers, tools_rows, col_widths=[1.4, 1.8, 3.8])

# ── SECTION 8: MVP Fast-Track Checklist ──────────────────────────────────────
add_heading(doc, '8. MVP Fast-Track Checklist')
add_body(doc, 'For time-critical MVP engagements, the following is the minimum viable document set. '
         'All items are required before any production launch. Items marked [POST] may be completed '
         'immediately after go-live but must be delivered within 2 weeks.')

checklist = [
    ('Strategy', [
        'Product Vision Doc',
        'Business Requirements Document (BRD) — even if lightweight (1–2 pages)',
    ]),
    ('Requirements', [
        'PRD (features scoped to MVP only)',
        'User Stories (MVP sprint backlog)',
        'API Specification (endpoints built for MVP)',
    ]),
    ('Architecture', [
        'Domain Model (bounded contexts + key aggregates)',
        'Database Schema (Prisma schema committed)',
        'System Architecture Doc (1-page architecture overview acceptable)',
    ]),
    ('Security', [
        'Security Architecture Doc (auth, encryption, secrets handling)',
        'DPIA [POST — within 2 weeks of launch if MVP handles PII]',
    ]),
    ('Dev Ops', [
        'Coding Standards Guide (can reference this Playbook)',
        'Git Strategy Doc (branching model defined)',
        'Test Plan (unit + integration coverage targets)',
        'Environment Setup Guide (README-level acceptable for MVP)',
    ]),
    ('Deployment', [
        'Deployment Runbook (step-by-step for first prod deploy)',
        'Monitoring & Alerting Plan (at minimum: uptime + error rate alerts active)',
        'Incident Response Playbook [POST — within 2 weeks]',
    ]),
    ('Handoff', [
        'API Reference Docs [POST — within 2 weeks]',
        'Developer Onboarding Guide [POST — within 2 weeks]',
    ]),
]

for category, items in checklist:
    add_heading(doc, category, 3)
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'[ ]  {item}')
        run.font.size = Pt(10)

doc.add_paragraph()
add_body(doc,
         'Total MVP doc count: 18 documents (12 pre-launch + 6 post-launch within 2 weeks). '
         'No MetronBase Labs project may go to production without completing the 12 pre-launch items above.',
         italic=True)

# ── SAVE ──────────────────────────────────────────────────────────────────────
output_path = r'c:\ReveenueCollect\docs\MetronBase_Labs_Playbook.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
